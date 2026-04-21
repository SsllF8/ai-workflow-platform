"""
文件上传节点 - 支持多种文件格式
"""
import os
import io
import base64
from nodes import BaseNode, NodeResult, register_node


def extract_pdf(file_bytes: bytes) -> str:
    """提取 PDF 文本内容"""
    import pdfplumber
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts) if text_parts else "未能提取到文本内容（可能是扫描件或纯图片PDF）"
    except Exception as e:
        return f"PDF 解析失败：{str(e)}"


def extract_docx(file_bytes: bytes) -> str:
    """提取 Word 文档文本"""
    from docx import Document
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 也提取表格
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))
        result = "\n".join(paragraphs)
        if tables_text:
            result += "\n\n[表格内容]\n" + "\n".join(tables_text)
        return result if result else "文档内容为空"
    except Exception as e:
        return f"Word 文档解析失败：{str(e)}"


def extract_xlsx(file_bytes: bytes) -> str:
    """提取 Excel 表格文本"""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        sheets_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rows_text.append(" | ".join(cells))
            if rows_text:
                sheets_text.append(f"[工作表: {sheet_name}]\n" + "\n".join(rows_text))
        wb.close()
        return "\n\n".join(sheets_text) if sheets_text else "Excel 内容为空"
    except Exception as e:
        return f"Excel 解析失败：{str(e)}"


def extract_text(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """提取纯文本文件"""
    try:
        text = file_bytes.decode(encoding)
        return text.strip() if text.strip() else "文件内容为空"
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("gbk")
            return text.strip() if text.strip() else "文件内容为空"
        except Exception:
            return "文件编码无法识别，请确保是文本文件"


def extract_image_info(file_bytes: bytes, filename: str) -> str:
    """提取图片基本信息（AI 节点可以用文件路径进一步处理）"""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        info = [
            f"[图片文件: {filename}]",
            f"尺寸: {img.width} x {img.height} 像素",
            f"格式: {img.format or '未知'}",
            f"模式: {img.mode}",
            f"文件大小: {len(file_bytes) / 1024:.1f} KB",
        ]
        # 如果有 EXIF 信息
        if hasattr(img, '_getexif') and img._getexif():
            info.append("包含 EXIF 元数据")
        return "\n".join(info)
    except Exception as e:
        return f"图片信息提取失败：{str(e)}"


def extract_csv(file_bytes: bytes, encoding: str = "utf-8") -> str:
    """提取 CSV 文件文本"""
    import csv
    try:
        text = file_bytes.decode(encoding)
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            rows.append(" | ".join(row))
        return "\n".join(rows) if rows else "CSV 内容为空"
    except UnicodeDecodeError:
        text = file_bytes.decode("gbk")
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            rows.append(" | ".join(row))
        return "\n".join(rows) if rows else "CSV 内容为空"


def extract_markdown(file_bytes: bytes) -> str:
    """提取 Markdown 文件"""
    try:
        text = file_bytes.decode("utf-8")
        return text.strip() if text.strip() else "文件内容为空"
    except Exception:
        return "Markdown 文件解析失败"


# 文件类型映射
EXTRACTORS = {
    ".pdf": ("PDF 文档", extract_pdf),
    ".docx": ("Word 文档", extract_docx),
    ".doc": ("Word 文档（旧版）", extract_docx),
    ".xlsx": ("Excel 表格", extract_xlsx),
    ".xls": ("Excel 表格（旧版）", extract_xlsx),
    ".txt": ("纯文本", extract_text),
    ".md": ("Markdown", extract_markdown),
    ".csv": ("CSV 表格", extract_csv),
    ".json": ("JSON", extract_text),
    ".xml": ("XML", extract_text),
    ".html": ("HTML", extract_text),
    ".htm": ("HTML", extract_text),
    ".log": ("日志文件", extract_text),
    ".py": ("Python 脚本", extract_text),
    ".java": ("Java 代码", extract_text),
    ".js": ("JavaScript 代码", extract_text),
    ".css": ("CSS 样式", extract_text),
    ".sql": ("SQL 脚本", extract_text),
    ".yaml": ("YAML", extract_text),
    ".yml": ("YAML", extract_text),
    ".toml": ("TOML", extract_text),
    ".ini": ("INI 配置", extract_text),
    ".conf": ("配置文件", extract_text),
    ".pptx": ("PowerPoint（仅文件名）", None),
    # 图片类
    ".png": ("图片", extract_image_info),
    ".jpg": ("图片", extract_image_info),
    ".jpeg": ("图片", extract_image_info),
    ".gif": ("图片", extract_image_info),
    ".bmp": ("图片", extract_image_info),
    ".webp": ("图片", extract_image_info),
    ".svg": ("SVG 图片", extract_text),
}


@register_node("file_upload")
class FileUploadNode(BaseNode):
    name = "文件上传"
    icon = "📎"
    description = "上传文件（PDF/Word/Excel/图片/文本等），提取内容传递给下游节点"
    config_fields = []  # 文件通过 st.file_uploader 上传，不需要额外配置

    # 接受的文件类型描述
    accept_types = [
        ".pdf", ".docx", ".doc", ".xlsx", ".xls",
        ".txt", ".md", ".csv", ".json", ".xml", ".html",
        ".py", ".java", ".js", ".sql", ".yaml", ".yml",
        ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg",
        ".pptx", ".log"
    ]

    def execute(self, input_data: str, config: dict) -> NodeResult:
        """
        执行文件解析。
        注意：实际文件内容由 app.py 在运行前写入 config["file_bytes"] 和 config["filename"]
        """
        file_bytes = config.get("file_bytes")
        filename = config.get("filename", "")

        if not file_bytes:
            # 尝试用 input_data（上游传入的文件路径或其他内容）
            if input_data:
                return NodeResult(success=True, output=input_data)
            return NodeResult(success=False, output="", error="请先上传文件")

        ext = os.path.splitext(filename.lower())[1]

        if ext not in EXTRACTORS:
            return NodeResult(
                success=False, output="",
                error=f"不支持的文件格式: {ext}\n支持的格式: {', '.join(sorted(set(EXTRACTORS.keys())))}"
            )

        file_type_name, extractor = EXTRACTORS[ext]

        if extractor is None:
            return NodeResult(
                success=True,
                output=f"[{file_type_name}] {filename}\n\n该格式暂不支持内容提取，仅支持文件名传递。"
            )

        try:
            if ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"]:
                content = extractor(file_bytes, filename)
            else:
                content = extractor(file_bytes)

            header = f"[{file_type_name}] {filename}\n{'=' * 40}\n\n"
            return NodeResult(success=True, output=header + content)

        except Exception as e:
            return NodeResult(success=False, output="", error=f"文件解析出错：{str(e)}")
